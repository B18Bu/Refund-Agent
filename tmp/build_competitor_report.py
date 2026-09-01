from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE


OUT = "docs/电商客诉舆情退赔决策系统竞品调研报告.docx"

# 版式：standard_business_brief；中文字体属于统一的 CJK 字体覆盖。
COLORS = {
    "navy": "203748",
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "muted": "5B6573",
    "pale_blue": "E8EEF5",
    "light": "F4F6F9",
    "gold": "7A5A00",
    "red": "9B1C1C",
    "green": "1F3A5F",
    "line": "D9E0E8",
}


def set_run_font(run, size=None, color=None, bold=None, italic=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = qn(f"w:{side}")
        node = margins.find(tag)
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="D9E0E8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def style_document(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    specs = {
        "Heading 1": (16, COLORS["blue"], 16, 8),
        "Heading 2": (13, COLORS["blue"], 12, 6),
        "Heading 3": (12, COLORS["dark_blue"], 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10
        style.paragraph_format.keep_with_next = True

    for name in ("Table Text", "Small Text", "Callout"):
        if name not in [s.name for s in doc.styles]:
            doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    table_text = doc.styles["Table Text"]
    table_text.base_style = normal
    table_text.font.size = Pt(9)
    table_text.paragraph_format.space_after = Pt(0)
    table_text.paragraph_format.line_spacing = 1.05
    small = doc.styles["Small Text"]
    small.base_style = normal
    small.font.size = Pt(9)
    small.font.color.rgb = RGBColor.from_string(COLORS["muted"])
    small.paragraph_format.space_after = Pt(3)
    callout = doc.styles["Callout"]
    callout.base_style = normal
    callout.font.size = Pt(11)
    callout.paragraph_format.space_after = Pt(0)
    callout.paragraph_format.line_spacing = 1.15

    header_p = sec.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header_p.add_run("电商客诉舆情退赔决策系统｜采购前行业调研")
    set_run_font(r, size=8.5, color=COLORS["muted"])
    footer_p = sec.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer_p.add_run("公开资料初筛版｜2026-09-01｜第 ")
    set_run_font(r, size=8.5, color=COLORS["muted"])
    add_page_field(footer_p)
    r = footer_p.add_run(" 页")
    set_run_font(r, size=8.5, color=COLORS["muted"])


def add_text(doc, text, style=None, align=None, size=None, color=None, bold=None, italic=None, before=None, after=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_lead(doc, label, text, color=None):
    p = doc.add_paragraph()
    r = p.add_run(label + "：")
    set_run_font(r, bold=True, color=color or COLORS["navy"])
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_source_note(doc, text):
    p = add_text(doc, text, style="Small Text", before=4, after=4)
    return p


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_borders(table)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, COLORS["pale_blue"])
        p = cell.paragraphs[0]
        p.style = "Table Text"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, size=9, color=COLORS["navy"], bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.style = "Table Text"
            if i in (0, 1):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(value)
            set_run_font(r, size=9)
    return table


def add_callout(doc, title, text, fill="F4F6F9"):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, COLORS["line"])
    shade(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    p.style = "Callout"
    r = p.add_run(title + "  ")
    set_run_font(r, bold=True, color=COLORS["navy"])
    r = p.add_run(text)
    set_run_font(r)
    return table


def add_vendor_card(doc, name, country, positioning, overlap, caveat):
    add_text(doc, name + "（" + country + "）", style="Heading 3")
    add_lead(doc, "公开定位", positioning)
    add_lead(doc, "与本项目的交集", overlap, COLORS["green"])
    add_lead(doc, "采购判断", caveat, COLORS["gold"])


def build_document():
    doc = Document()
    style_document(doc)

    # 封面：editorial_cover 的正文适配版；标题和上下留白为唯一 named override。
    add_text(doc, "采购前行业调研", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color=COLORS["gold"], bold=True, before=112, after=18)
    add_text(doc, "电商客诉舆情退赔\n决策系统竞品调研报告", align=WD_ALIGN_PARAGRAPH.CENTER, size=28, color=COLORS["navy"], bold=True, after=12)
    add_text(doc, "退款欺诈与滥用风控、售后自动化、舆情感知的市场格局与采购建议", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, color=COLORS["muted"], after=42)
    add_text(doc, "适用场景：项目立项前的商业与采购调研", align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, color=COLORS["muted"], italic=True, after=8)
    add_text(doc, "资料截止：2026 年 9 月 1 日｜版本：公开资料初筛版", align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, color=COLORS["muted"], after=18)
    add_text(doc, "重要说明：本报告仅基于公开官网与公开产品资料梳理；未公开的功能、价格、中文本地化、数据处理方式与交付能力均标注为“需 PoC/RFI 验证”，不构成采购承诺。", align=WD_ALIGN_PARAGRAPH.CENTER, size=9, color=COLORS["muted"], after=0)
    doc.add_page_break()

    add_text(doc, "执行摘要", style="Heading 1")
    add_callout(doc, "结论", "目标系统并非单一成熟品类，而是“退款/退赔决策编排层”。海外反欺诈平台在账户、支付、拒付与策略滥用识别上重叠最高；国内厂商以业务风控、内容安全和舆情监测等组件型能力为主。采购应避免把任一单品等同于完整方案。", "E8EEF5")
    add_lead(doc, "建议的采购路径", "优先采用“自建审批与决策编排核心 + 采购可插拔风控/舆情能力”的组合式方案。这样既保留订单金额、证据、人工审批、幂等与审计闭环，也能在后续更换模型或供应商时降低迁移成本。")
    add_lead(doc, "最接近的海外标杆", "Forter、Riskified、Signifyd、Sift 和 ClearSale。它们重点解决电商欺诈、拒付、账户攻击或策略滥用，其中退款滥用通常属于“policy abuse”或“post-purchase abuse”能力，而非包含 OCR、舆情和人工审批状态机的完整售后工作台。")
    add_lead(doc, "中国市场的关键候选", "数美、同盾、百融可作为风控能力候选；蚁坊、识微可作为舆情感知候选；聚水潭、旺店通可作为订单与售后流程集成参照。中国平台的内建售后能力很强，但一般不对外提供可独立采购的全链路风控产品。")
    add_lead(doc, "项目的可辩护差异", "以订单、退款描述、凭证图片/OCR、风险信号与舆情信号为输入，由确定性红线和可解释规则控制自动退赔，并在不确定情形进入可恢复的人工审批闭环。公开资料中未能确认单一厂商同时覆盖这一完整链路。")
    add_lead(doc, "采购红线", "自动退赔不能被外部风控评分直接触发；金额阈值、硬性证据规则、审批权限、幂等控制与最终支付动作必须保留在本系统的确定性控制面。")
    doc.add_page_break()

    add_text(doc, "调研边界与市场定义", style="Heading 1")
    add_lead(doc, "目标问题", "当电商平台在促销、质量事故、物流异常或公共事件中出现集中退款时，在降低人工审核负担的同时，识别刷单、虚假凭证、恶意退款和策略套利，并让每次决定可复核、可追溯。")
    add_lead(doc, "纳入规则", "优先纳入公开资料显示与电商欺诈、退款/退货滥用、售后流程、内容风控或舆情监测有关的产品；同时纳入支付风控和 ERP/OMS 等相邻产品，避免误把“局部能力”写成“直接竞品”。")
    add_lead(doc, "不纳入规则", "不把平台内部能力、咨询服务、纯客服工单工具或泛用大模型当作可直接替换方案。它们可以是能力标杆或集成对象，但不应进入同一采购评分表。")
    add_source_note(doc, "分类口径：直接竞品＝可对退款欺诈/策略滥用做决策或保护；组件型竞品＝只覆盖风控、舆情、售后或支付的一段能力；生态基准＝平台内建能力或订单系统。")
    add_matrix(doc, ["层级", "采购对象", "解决的问题", "与目标系统的关系"], [
        ["直接竞品", "Forter、Riskified、Signifyd、Sift、ClearSale", "欺诈、拒付、账户攻击、策略滥用", "风险决策能力重叠最高；需补售后编排与中国本地化"],
        ["组件型竞品", "数美、同盾、百融、Stripe Radar、Adyen", "业务/支付风控与策略规则", "可替换或增强风险评分能力"],
        ["感知与流程", "蚁坊、识微、聚水潭、旺店通、Loop/Narvar", "舆情、订单、售后与退货体验", "补足输入和流程，不替代最终决策"],
        ["生态基准", "淘宝/天猫、京东、拼多多等平台能力", "平台内售后与治理", "体验、规则与效率基准；通常不可独立采购"],
    ], [1500, 2350, 2750, 2760])
    doc.add_page_break()

    add_text(doc, "竞品全景与初筛矩阵", style="Heading 1")
    add_source_note(doc, "说明：“高/中/低”表示与目标系统的公开能力重叠程度，并非厂商实力排名；“需验证”表示官网公开资料不足以确认。")
    add_matrix(doc, ["厂商/产品", "国家或主要市场", "定位", "重叠度", "初筛结论"], [
        ["Forter", "以色列起源/全球", "数字商业信任与欺诈防护", "高", "海外能力标杆；重点验证退款策略滥用与中国部署"],
        ["Riskified", "以色列起源/全球", "电商风险、拒付与策略保护", "高", "海外能力标杆；重点验证售后与证据链覆盖"],
        ["Signifyd", "美国", "Commerce Protection", "高", "适合比较拒付保障与滥用防护模式"],
        ["Sift", "美国", "Fraud Prevention Platform", "高", "适合比较账户、支付与滥用识别"],
        ["ClearSale", "巴西/全球", "反欺诈与拒付保护", "中高", "适合比较人工审核与拒付保护运营"],
        ["数美科技", "中国", "内容安全与智能业务风控", "中高", "国内风险与内容能力候选；退款专用流程需验证"],
        ["同盾", "中国", "智能分析、欺诈风险预测与决策", "中高", "国内风险决策候选；证据和舆情能力需验证"],
        ["百融智能", "中国", "数据智能与风险决策", "中", "国内数据/风控候选；电商退款场景需验证"],
        ["蚁坊软件", "中国", "全网舆情监测与分析", "中", "舆情感知组件候选，不替代交易风控"],
        ["识微科技", "中国", "舆情监测与预警", "中", "舆情感知组件候选，不替代交易风控"],
        ["聚水潭/旺店通", "中国", "电商 ERP、订单与售后流程", "低", "订单/售后集成参照，不是反欺诈产品"],
    ], [2050, 1500, 2200, 700, 2910])
    add_source_note(doc, "采购提示：海外产品的“退款滥用”能力常以 policy abuse、promotion abuse、returns abuse 或 post-purchase abuse 命名；RFI 时必须要求厂商按本项目案例演示，而不能只接受支付拒付指标。")
    doc.add_page_break()

    add_text(doc, "海外直接竞品：能力标杆", style="Heading 1")
    add_vendor_card(doc, "Forter", "以色列起源/全球", "官网将产品定位为数字商业场景的信任与智能决策平台，面向欺诈、账户与交易风险。", "可对比身份/行为信号汇聚、实时决策、策略滥用防护和商家运营模型。", "公开首页不足以证明其覆盖中文 OCR、公开舆情或中国电商退款审批；要求演示退款滥用、售后证据与回退逻辑。")
    add_vendor_card(doc, "Riskified", "以色列起源/全球", "以电商风险管理、拒付保护、账户安全和策略保护为主要公开方向。", "可对比“商家承担损失”与“平台给出风险判断”两类商业模式，以及退货/退款政策滥用检测思路。", "其官网受到访问保护，功能边界需通过销售资料和 PoC 二次核验；不可仅据市场口碑认定全部适配。")
    add_vendor_card(doc, "Signifyd", "美国", "以 Commerce Protection 为主张，围绕订单欺诈、拒付和滥用风险提供商业保护。", "适合研究订单级风险决策、拒付保障和审批队列运营。", "需验证是否能将订单描述、图片、OCR 等作为特征输入，及中国数据处理与合规安排。")
    add_vendor_card(doc, "Sift", "美国", "官网定位为 Fraud Prevention Platform，覆盖支付欺诈、账户防护和滥用识别。", "适合比较跨账号、设备、行为关联和实时评分能力。", "并非售后工作流产品；需通过 API 能力判断能否嵌入本系统而非替换系统。")
    add_vendor_card(doc, "ClearSale", "巴西/全球", "公开资料强调电商欺诈预防与拒付保护。", "适合比较“自动判定 + 人工审核”的运营分工。", "需验证中国市场交付、语言和证据图片处理能力。")
    doc.add_page_break()

    add_text(doc, "国内与相邻能力候选", style="Heading 1")
    add_vendor_card(doc, "数美科技", "中国", "官网公开定位为内容安全与智能业务风控。", "其“内容 + 业务风险”双线能力，与退款描述、图片审核、恶意行为识别的输入结构最相近。", "必须确认是否提供退款/售后欺诈模型、图像凭证一致性、可解释输出、私有化或数据驻留与审计接口。")
    add_vendor_card(doc, "同盾", "中国", "官网公开描述聚焦智能分析与决策，以及欺诈风险预测。", "适合补充身份、设备、行为和关联网络风险信号，并作为规则/模型评分的外部输入。", "公开资料不能确认其具备舆情与凭证 OCR 的全链路能力；更适合作为风险引擎候选。")
    add_vendor_card(doc, "百融智能", "中国", "官网公开定位为数据智能服务。", "可作为风险数据与决策能力的中国候选之一，适合在高风险用户、异常行为和数据覆盖维度做 RFI 对比。", "必须谨慎验证电商退款专用场景、数据来源合法性、可解释性、保留期及个人信息处理边界。")
    add_vendor_card(doc, "蚁坊软件与识微科技", "中国", "两者公开定位均涵盖全网舆情监测、分析、预警或响应。", "可为集中客诉、商品质量事故、热点扩散等提供宏观风险信号，帮助动态收紧或放宽人工审核队列。", "舆情不能单独判定单笔退款真假；应只作为群体性风险与运营优先级输入。")
    add_vendor_card(doc, "聚水潭与旺店通", "中国", "公开市场定位为电商 ERP/OMS、订单与售后流程管理。", "可作为订单状态、物流、履约和售后工单的数据来源及集成参照。", "它们不是反欺诈替代品；应评估 API、事件回调、渠道覆盖和数据一致性。")
    doc.add_page_break()

    add_text(doc, "能力对标：目标系统的不可替代部分", style="Heading 1")
    add_lead(doc, "订单与金额", "成熟风控平台普遍擅长交易级风险，但本项目需要把金额、履约阶段、退款频率、优惠活动和品类规则放进可审计的同一决策上下文。")
    add_lead(doc, "证据与多模态", "用户描述、破损照片、快递单、聊天截图都是不可信输入。OCR 置信度、图文一致性、篡改迹象和缺失材料应当触发确定性降级，而不是由语言模型单独“判断真实”。")
    add_lead(doc, "舆情感知", "舆情系统擅长发现群体性事件；它能解释“为什么某批订单应优先人工审核”，却不能替代订单级欺诈证明。最佳实践是把舆情等级作为阈值、队列优先级或审核资源的调节信号。")
    add_lead(doc, "人工审批与状态恢复", "采购的风控评分必须进入受控工作流：暂停、分配、审批、拒绝、恢复和审计。自动退赔与实际支付接口应分离，避免把系统状态写成真实资金动作。")
    add_lead(doc, "可解释与可复核", "每次决定至少保留：命中的硬规则、外部评分与版本、OCR 置信度、舆情信号、审批人和时间线。评分供应商无法提供可解释字段时，不应拥有自动放行权。")
    add_callout(doc, "因此", "最稳妥的定位不是“替代风控厂商”，而是作为电商退款场景的决策编排与审计控制面：向下集成风险、OCR、舆情和订单系统，向上提供自动处理与人工复核。", "F4F6F9")
    doc.add_page_break()

    add_text(doc, "采购选型建议", style="Heading 1")
    add_lead(doc, "方案 A：单一海外平台", "适用于跨境电商、支付欺诈和拒付损失是主要问题、且可接受数据跨境与海外交付的组织。优点是模型和运营成熟；缺点是中文证据、国内舆情、国产平台集成和本地合规常需二次开发。")
    add_lead(doc, "方案 B：国内风控 + 舆情组件 + 自建编排层（推荐）", "适用于以中国电商渠道为主、需要把退款审核嵌入既有后台的组织。风险引擎、舆情与 OCR 均可替换；核心业务规则、审批、幂等和审计留在本系统。")
    add_lead(doc, "方案 C：完全自建", "适用于数据规模大、反作弊特征差异显著、且已有算法/运营团队的组织。优点是控制力最强；缺点是冷启动数据、模型维护和对抗升级成本高，不建议在 MVP 阶段承担全部责任。")
    add_matrix(doc, ["评价维度", "A 单一海外平台", "B 组合式方案（推荐）", "C 完全自建"], [
        ["上线速度", "快", "中", "慢"],
        ["中国电商与中文证据适配", "需验证", "较强", "可定制"],
        ["退款工作流控制力", "中", "高", "高"],
        ["供应商锁定风险", "高", "中", "低"],
        ["模型/运营成本", "订阅与服务费", "集成与多供应商管理", "长期研发与运营"],
        ["建议", "作为标杆或跨境候选", "优先进入 RFI/PoC", "仅作为中长期路线"],
    ], [1800, 2520, 2520, 2520])
    doc.add_page_break()

    add_text(doc, "RFI（需求信息征集）问题清单", style="Heading 1")
    add_source_note(doc, "以下问题用于把“宣传能力”转成可验证的采购条款。建议要求厂商逐项标明：标准能力、配置能力、二开能力或不支持。")
    add_matrix(doc, ["主题", "必须回答的问题", "验收证据"], [
        ["退款/退货滥用", "能否区分真实售后、策略套利、恶意退款、团伙刷单？是否已有电商退款案例？", "按给定工单样本演示命中原因与复核路径"],
        ["多模态证据", "是否支持文本、图片、OCR 结构化字段和图文一致性？模型失败如何降级？", "错误/低置信度样本的输出与回退日志"],
        ["规则与解释", "是否支持金额、履约、频率、黑白名单等确定性红线？每次评分能否返回原因码、特征摘要和模型版本？", "接口响应、规则版本与审计记录"],
        ["人工审核", "是否支持队列、优先级、复核反馈、双人审批和审计？", "端到端审批演示与权限模型说明"],
        ["中国合规", "数据处理地点、子处理商、留存/删除、脱敏、跨境路径和安全认证为何？", "数据处理协议、架构图、合规清单"],
        ["集成与可靠性", "是否有 REST/Webhook、幂等键、超时策略、限流、SLA 和故障回退？", "API 文档、压测数据、故障演练记录"],
        ["商业模式", "按订单、按评分、按保障金额还是按年订阅计费？误判、拒付和服务边界如何定义？", "报价单、容量阶梯、退出与数据导出条款"],
    ], [1450, 5000, 2910])
    doc.add_page_break()

    add_text(doc, "PoC（概念验证）设计", style="Heading 1")
    add_lead(doc, "样本设计", "使用脱敏历史工单或合成样本，至少覆盖：正常退款、证据不足、疑似恶意退款、集中客诉事件、重复提交、图片低质量和 OCR 失败。样本标签由业务与风控共同复核，避免把历史人工决定直接当作真值。")
    add_lead(doc, "比较方式", "同一批样本同时送入候选供应商与本系统基线规则；记录自动放行率、人工转入率、误放行、误拦截、原因码完整率、接口时延、失败回退和审核人员复核耗时。")
    add_lead(doc, "硬性门槛", "所有硬规则命中时必须 100% 转人工或拒绝；模型超时、非法输出、证据低置信度时必须保守降级；每条自动决定必须能关联到规则/评分版本和输入摘要。")
    add_lead(doc, "通过标准", "不预设脱离业务基线的“准确率神话”。候选方应在相同安全门槛下，用可复算的结果证明：减少可安全自动处理的人工量、不过度增加误伤、并完整输出可审计理由。")
    add_lead(doc, "安全约束", "PoC 不接真实支付接口；输入图片和文本先脱敏；使用独立测试租户与最小权限 API Key；删除策略、日志脱敏和数据回收写入 PoC 协议。")
    add_callout(doc, "推荐决策门", "先通过“安全与解释性”门，再比较效率。任何无法说明为什么自动放行、无法隔离数据或无法在异常时保守回退的产品，即使平均准确率较高，也不应进入生产自动退赔链路。", "F4F6F9")
    doc.add_page_break()

    add_text(doc, "建议的目标架构与推进节奏", style="Heading 1")
    add_lead(doc, "控制面", "订单服务、退款规则、审批权限、审批锁、幂等键、状态机与审计日志由本系统掌控。它们决定是否允许自动退赔，不能外包给黑箱评分。")
    add_lead(doc, "能力面", "OCR、外部风控评分、舆情信号、图像安全、设备/账号风险等以异步或可超时调用的插件接入。单个供应商失败不得阻塞其他结果，且异常默认转人工。")
    add_lead(doc, "数据面", "输入材料最小化、字段分级、脱敏传输、保存期限可配置。舆情只保留必要的聚合信号和来源摘要，不将原始投诉内容作为不受控提示词。")
    add_lead(doc, "第 1 阶段", "先完成 2 至 3 家国内组件和 1 家海外标杆的 RFI；明确数据本地化、场景覆盖和商务边界。")
    add_lead(doc, "第 2 阶段", "选择 2 家进入隔离 PoC，以同一套退款样本和安全门槛运行；业务、风控、法务与技术共同签字。")
    add_lead(doc, "第 3 阶段", "先灰度到“只给风险建议 + 人工审批”，通过稳定性、误伤率和审计抽检后，再把满足全部低风险条件的小额工单开放为自动处理。")
    doc.add_page_break()

    add_text(doc, "公开资料来源与使用说明", style="Heading 1")
    add_lead(doc, "检索时间", "2026 年 9 月 1 日。来源优先采用厂商官网的产品主页；由于部分官网受动态渲染、地区跳转或反爬保护影响，本报告不把未访问到的页面视为功能否定。")
    sources = [
        ["Forter", "https://www.forter.com/", "官网首页：数字商业信任与智能决策定位"],
        ["Riskified", "https://www.riskified.com/", "官网主页与解决方案页面：电商风险与策略保护（访问保护，需 RFI 核验）"],
        ["Signifyd", "https://www.signifyd.com/", "官网主页：Commerce Protection 定位"],
        ["Sift", "https://sift.com/products/", "产品页：Fraud Prevention Platform"],
        ["ClearSale", "https://clear.sale/", "官网主页：反欺诈与拒付保护"],
        ["Stripe Radar", "https://stripe.com/radar", "产品页：AI-powered Fraud Detection"],
        ["Adyen RevenueProtect", "https://www.adyen.com/risk-management", "风险管理产品资料"],
        ["数美科技", "https://www.ishumei.com/", "官网首页：内容安全与智能业务风控"],
        ["同盾", "https://www.tongdun.cn/", "官网首页：智能分析、决策与欺诈风险预测"],
        ["百融智能", "https://www.brgroup.com/", "官网主页：数据智能服务"],
        ["蚁坊软件", "https://www.eefung.com/", "官网首页：全网舆情监测与分析"],
        ["识微科技", "https://www.ciswei.com/", "官网公开资料：舆情监测与预警"],
        ["聚水潭", "https://www.jushuitan.com/", "官网公开资料：电商业务数字化与订单流程"],
        ["旺店通", "https://www.wangdian.cn/", "官网公开资料：电商 ERP/订单与售后流程"],
        ["Loop Returns", "https://www.loopreturns.com/", "退货与换货体验的相邻国际参考"],
        ["Narvar", "https://corp.narvar.com/", "退货、履约与客户体验的相邻国际参考"],
    ]
    add_matrix(doc, ["来源", "链接", "本报告中的用途"], sources, [1650, 4100, 3610])
    add_source_note(doc, "引用方式：本报告仅据公开资料描述公开定位，不推断未披露功能。采购决策应以签署后的产品说明、RFI 回复、PoC 结果、数据处理协议和 SLA 为准。")

    doc.core_properties.title = "电商客诉舆情退赔决策系统竞品调研报告"
    doc.core_properties.subject = "采购前行业与竞品调研"
    doc.core_properties.author = "项目调研组"
    doc.core_properties.comments = "公开资料初筛版"
    doc.save(OUT)


if __name__ == "__main__":
    build_document()
